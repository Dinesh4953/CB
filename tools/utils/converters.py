import os
import io
import tempfile
from PIL import Image
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import pypdf
from pdf2image import convert_from_bytes, convert_from_path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pytesseract
import docx2txt
import mimetypes

class TextToPdfConverter:
    """Convert plain text to PDF"""
    
    @staticmethod
    def convert(text_content, filename="document.pdf", font_size=12, title=None):
        """
        Convert text to PDF
        
        Args:
            text_content: String content to convert
            filename: Output filename
            font_size: Font size (default 12)
            title: Optional title for the PDF
            
        Returns:
            BytesIO object with PDF content
        """
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        elements = []
        
        styles = getSampleStyleSheet()
        
        # Add title if provided
        if title:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='#1f4788',
                spaceAfter=30,
                alignment=1  # Center
            )
            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 0.3*inch))
        
        # Add content
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['BodyText'],
            fontSize=font_size,
            leading=font_size + 4,
            spaceAfter=12
        )
        
        # Split text into paragraphs
        paragraphs = text_content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                elements.append(Paragraph(para.strip(), content_style))
                elements.append(Spacer(1, 0.2*inch))
        
        doc.build(elements)
        pdf_buffer.seek(0)
        return pdf_buffer

class PdfToTextConverter:
    """Convert PDF to plain text"""
    
    @staticmethod
    def convert(pdf_file):
        """
        Extract text from PDF
        
        Args:
            pdf_file: PDF file bytes or file path
            
        Returns:
            String with extracted text
        """
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        
        for page_num, page in enumerate(reader.pages):
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.extract_text()
        
        return text

class WordToPdfConverter:
    """Convert Word (DOCX) documents to PDF"""
    
    @staticmethod
    def convert(docx_file, filename="document.pdf"):
        """
        Convert DOCX to PDF
        
        Args:
            docx_file: DOCX file bytes or file path
            filename: Output filename
            
        Returns:
            BytesIO object with PDF content
        """
        # Extract text from DOCX
        if isinstance(docx_file, bytes):
            docx_buffer = io.BytesIO(docx_file)
            doc = Document(docx_buffer)
        else:
            doc = Document(docx_file)
        
        # Create PDF
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                elements.append(Paragraph(paragraph.text, styles['BodyText']))
                elements.append(Spacer(1, 0.1*inch))
        
        # Add tables if they exist
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), '#cccccc'),
                    ('TEXTCOLOR', (0, 0), (-1, 0), '#000000'),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), '#ffffff'),
                    ('GRID', (0, 0), (-1, -1), 1, '#cccccc'),
                ]))
                elements.append(t)
                elements.append(Spacer(1, 0.2*inch))
        
        pdf_doc.build(elements)
        pdf_buffer.seek(0)
        return pdf_buffer

class PdfToWordConverter:
    """Convert PDF to Word (DOCX) documents"""
    
    @staticmethod
    def convert(pdf_file, filename="document.docx"):
        """
        Convert PDF to DOCX
        
        Args:
            pdf_file: PDF file bytes or file path
            filename: Output filename
            
        Returns:
            BytesIO object with DOCX content
        """
        # Extract text from PDF
        if isinstance(pdf_file, bytes):
            reader = pypdf.PdfReader(io.BytesIO(pdf_file))
        else:
            reader = pypdf.PdfReader(pdf_file)
        
        # Create DOCX document
        doc = Document()
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            
            if page_num > 0:
                doc.add_page_break()
            
            if text.strip():
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
        
        # Save to BytesIO
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer

class ImageToPdfConverter:
    """Convert images to PDF"""
    
    @staticmethod
    def convert(image_files, filename="document.pdf", orientation='portrait'):
        """
        Convert one or multiple images to PDF
        
        Args:
            image_files: List of image files or single file
            filename: Output filename
            orientation: 'portrait' or 'landscape'
            
        Returns:
            BytesIO object with PDF content
        """
        if not isinstance(image_files, list):
            image_files = [image_files]
        
        images = []
        for img_file in image_files:
            if isinstance(img_file, bytes):
                img = Image.open(io.BytesIO(img_file))
            else:
                img = Image.open(img_file)
            
            # Convert RGBA to RGB if necessary
            if img.mode in ('RGBA', 'LA', 'P'):
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = rgb_img
            
            # Resize based on orientation
            if orientation == 'landscape':
                img = img.rotate(0)  # Keep as is
            
            images.append(img)
        
        pdf_buffer = io.BytesIO()
        images[0].save(pdf_buffer, format='PDF', save_all=True, append_images=images[1:], duration=100, loop=0)
        pdf_buffer.seek(0)
        return pdf_buffer

class PdfToImageConverter:
    """Convert PDF pages to images"""

    @staticmethod
    def convert(pdf_file, format='PNG', dpi=200, page_numbers=None):

        """
        Convert PDF pages to images

        Args:
            pdf_file: Uploaded PDF file
            format: PNG / JPG / BMP
            dpi: Image quality
            page_numbers: Optional page numbers

        Returns:
            List of image buffers
        """

        # Read uploaded PDF bytes
        pdf_bytes = pdf_file.read()

        # Convert PDF bytes into PIL images
        images = convert_from_bytes(
            pdf_bytes,
            dpi=dpi,
            first_page=page_numbers[0] if page_numbers else None,
            last_page=page_numbers[-1] if page_numbers else None
        )

        image_buffers = []

        for img in images:

            img_buffer = io.BytesIO()

            img.save(img_buffer, format=format)

            img_buffer.seek(0)

            image_buffers.append(img_buffer)

        return image_buffers

class ImageToTextConverter:
    """Extract text from images using OCR"""
    
    @staticmethod
    def convert(image_file, language='eng'):
        """
        Extract text from image using OCR
        
        Args:
            image_file: Image file bytes or file path
            language: OCR language (default 'eng')
            
        Returns:
            Extracted text string
        """
        if isinstance(image_file, bytes):
            img = Image.open(io.BytesIO(image_file))
        else:
            img = Image.open(image_file)
        
        text = pytesseract.image_to_string(img, lang=language)
        return text

class PdfMergeConverter:
    """Merge multiple PDFs into one"""

    @staticmethod
    def merge(pdf_files, filename="merged.pdf"):

        """
        Merge multiple PDF files

        Args:
            pdf_files: List of uploaded PDF files
            filename: Output filename

        Returns:
            BytesIO object with merged PDF
        """

        writer = pypdf.PdfWriter()

        for pdf_file in pdf_files:

            reader = pypdf.PdfReader(pdf_file)

            for page in reader.pages:
                writer.add_page(page)

        pdf_buffer = io.BytesIO()

        writer.write(pdf_buffer)

        pdf_buffer.seek(0)

        return pdf_buffer

class FileTypeDetector:
    """Detect file types"""
    
    @staticmethod
    def get_file_type(file_path):
        """Get MIME type of file"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type
    
    @staticmethod
    def is_pdf(file_path):
        return FileTypeDetector.get_file_type(file_path) == 'application/pdf'
    
    @staticmethod
    def is_image(file_path):
        mime_type = FileTypeDetector.get_file_type(file_path)
        return mime_type and mime_type.startswith('image/')
    
    @staticmethod
    def is_docx(file_path):
        return FileTypeDetector.get_file_type(file_path) == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
